"""
Shared fixtures for all tests.
Uses tmp_path for all file I/O — no pollution of real data directories.
"""
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch
from neural_search.ingestion.chunker import Chunk


# ── Sample data ───────────────────────────────────────────────────────────────
@pytest.fixture
def sample_chunks() -> list[Chunk]:
    return [
        Chunk(
            chunk_id=f"chunk_{i:04d}",
            doc_id="test_doc",
            source_file="test.pdf",
            page=i + 1,
            chunk_index=i,
            text=f"This is sample text for chunk number {i}. It contains meaningful content about topic {i}.",
            token_count=15 + i,
        )
        for i in range(10)
    ]


@pytest.fixture
def sample_pdf(tmp_path) -> Path:
    """Creates a minimal real PDF using pymupdf for parser tests."""
    try:
        import fitz
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((50, 50), f"Page {i+1} content. This is test document text for neural search.")
        path = tmp_path / "sample.pdf"
        doc.save(str(path))
        doc.close()
        return path
    except ImportError:
        pytest.skip("pymupdf not installed")


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    """Creates a minimal real DOCX for parser tests."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This is the introduction section with test content.")
        doc.add_heading("Methods", level=1)
        doc.add_paragraph("This section describes the methods used in this document.")
        path = tmp_path / "sample.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture
def mock_settings(tmp_path, monkeypatch):
    """Patch settings to use tmp_path for all data dirs."""
    from neural_search import config
    mock = MagicMock()
    mock.data_dir = tmp_path / "data"
    mock.qdrant_path = tmp_path / "data" / "qdrant"
    mock.bm25_index_path = tmp_path / "data" / "bm25_index"
    mock.embedding_model = "all-MiniLM-L6-v2"
    mock.qdrant_collection = "test_collection"
    mock.chunk_size = 256
    mock.chunk_overlap = 32
    mock.top_k = 5
    mock.rrf_k = 60
    mock.bm25_path_for = lambda slug: tmp_path / "data" / "bm25_index" / slug
    mock.documents_path_for = lambda slug: tmp_path / "data" / "documents" / slug
    mock.ensure_dirs = lambda: None
    monkeypatch.setattr(config, "settings", mock)
    return mock
